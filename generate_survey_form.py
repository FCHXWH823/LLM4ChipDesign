#!/usr/bin/env python3
"""
LLM4ChipDesign Course Prerequisite Survey Form Generator
This script generates a Word document survey form for students interested in the LLM4ChipDesign course.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def create_survey_form():
    """Create a comprehensive survey form for LLM4ChipDesign course prerequisites."""
    
    # Create a new document
    doc = Document()
    
    # Add document title
    title = doc.add_heading('LLM4ChipDesign Course Prerequisite Survey', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_heading('Student Background Assessment Form', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction paragraph
    intro = doc.add_paragraph()
    intro.add_run('Dear Student,\n\n').bold = True
    intro.add_run(
        'Welcome to the LLM4ChipDesign course! This course explores the intersection of '
        'Large Language Models (LLMs) and chip design, covering topics such as automated '
        'Verilog generation, testbench creation, SystemVerilog assertions, and hardware-software '
        'co-design using AI.\n\n'
        'To ensure you have the necessary background knowledge and to tailor the course content '
        'to your experience level, please complete this prerequisite assessment survey. '
        'Your responses will help us understand your current knowledge and provide appropriate '
        'support during the course.\n\n'
        'Please answer all questions honestly. This survey is for assessment purposes only '
        'and will not affect your grade.\n'
    )
    
    # Student Information Section
    doc.add_heading('Student Information', level=1)
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_fields = [
        ('Name:', ''),
        ('Student ID:', ''),
        ('Email:', ''),
        ('Academic Year/Level:', ''),
        ('Major/Program:', ''),
        ('Date:', '')
    ]
    
    for i, (field, value) in enumerate(info_fields):
        info_table.cell(i, 0).text = field
        info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        info_table.cell(i, 1).text = value
        
    doc.add_paragraph()  # Add space
    
    # Prerequisite Course Assessment
    doc.add_heading('Prerequisite Course Assessment', level=1)
    
    doc.add_paragraph(
        'Please indicate your experience level with the following subject areas '
        'that are fundamental to this course. Rate your experience using the scale:\n'
        '• No Experience (0): Never studied or worked with this area\n'
        '• Basic (1): Introductory course or minimal exposure\n'
        '• Intermediate (2): One or more courses with practical experience\n'
        '• Advanced (3): Extensive coursework and/or professional experience\n'
        '• Expert (4): Teaching/research level knowledge\n'
    )
    
    # Create prerequisite courses table
    prereq_table = doc.add_table(rows=1, cols=6)
    prereq_table.style = 'Table Grid'
    
    # Table headers
    headers = ['Course Area', 'No Exp (0)', 'Basic (1)', 'Inter (2)', 'Adv (3)', 'Expert (4)']
    for i, header in enumerate(headers):
        cell = prereq_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        if i == 0:
            cell.width = Inches(3.0)
        else:
            cell.width = Inches(0.8)
    
    # Prerequisite course areas based on repository content
    course_areas = [
        'Digital Logic Design & Boolean Algebra',
        'Computer Architecture & Organization',
        'Hardware Description Languages (Verilog/SystemVerilog)',
        'Programming Languages (Python, C/C++)',
        'Data Structures & Algorithms',
        'Machine Learning & Artificial Intelligence',
        'Natural Language Processing (NLP)',
        'Formal Verification & Model Checking',
        'Electronic Design Automation (EDA) Tools',
        'FPGA/ASIC Design Flow',
        'High-Level Synthesis (HLS)',
        'Hardware-Software Co-design',
        'Computer Systems & Operating Systems',
        'Software Engineering & Version Control',
        'Mathematics (Linear Algebra, Statistics)'
    ]
    
    # Add rows for each course area
    for course in course_areas:
        row = prereq_table.add_row()
        row.cells[0].text = course
        # Add checkbox-like symbols for rating
        for i in range(1, 6):
            row.cells[i].text = '☐'
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Add space
    
    # Specific Course Experience Section
    doc.add_heading('Specific Course Experience', level=1)
    
    doc.add_paragraph(
        'Please list any specific courses you have completed in the following areas '
        '(include course names/codes if possible):'
    )
    
    course_exp_table = doc.add_table(rows=8, cols=2)
    course_exp_table.style = 'Table Grid'
    
    specific_areas = [
        'Digital Circuit Design / Logic Design:',
        'Computer Architecture:',
        'Verilog/VHDL/SystemVerilog:',
        'Machine Learning / AI:',
        'Programming (Python/C/C++):',
        'Formal Methods / Verification:',
        'Embedded Systems:',
        'Other Relevant Courses:'
    ]
    
    for i, area in enumerate(specific_areas):
        course_exp_table.cell(i, 0).text = area
        course_exp_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        course_exp_table.cell(i, 0).width = Inches(2.5)
        course_exp_table.cell(i, 1).width = Inches(4.0)
    
    doc.add_paragraph()  # Add space
    
    # Tools and Software Experience
    doc.add_heading('Tools and Software Experience', level=1)
    
    doc.add_paragraph(
        'Please indicate your familiarity with the following tools and software '
        '(check all that apply):'
    )
    
    tools_table = doc.add_table(rows=1, cols=3)
    tools_table.style = 'Table Grid'
    
    # Create three columns for tools
    tools_categories = [
        'EDA Tools',
        'Programming Tools',
        'AI/ML Frameworks'
    ]
    
    tools_lists = [
        ['☐ Vivado (Xilinx)', '☐ Quartus (Intel)', '☐ ModelSim/QuestaSim', 
         '☐ Synopsys Tools', '☐ Cadence Tools', '☐ Verilator', '☐ GTKWave'],
        ['☐ Python', '☐ C/C++', '☐ MATLAB', '☐ Git/GitHub', 
         '☐ Jupyter Notebooks', '☐ Linux/Unix', '☐ Command Line'],
        ['☐ TensorFlow', '☐ PyTorch', '☐ OpenAI API', '☐ Hugging Face', 
         '☐ scikit-learn', '☐ NLTK/spaCy', '☐ Google Colab']
    ]
    
    for i, category in enumerate(tools_categories):
        tools_table.cell(0, i).text = category
        tools_table.cell(0, i).paragraphs[0].runs[0].bold = True
        
        # Add tools to each category
        for j, tool in enumerate(tools_lists[i]):
            if j == 0:
                tools_table.cell(0, i).text = f"{category}\n{tool}"
            else:
                tools_table.cell(0, i).text += f"\n{tool}"
    
    doc.add_paragraph()  # Add space
    
    # Project Experience Section
    doc.add_heading('Project Experience', level=1)
    
    project_questions = [
        'Have you worked on any hardware design projects? If yes, please describe briefly:',
        'Have you used any AI/ML tools or APIs in your projects? If yes, please describe:',
        'Have you written Verilog or VHDL code before? What was the complexity level?',
        'Have you used any LLMs (ChatGPT, Claude, etc.) for coding assistance? Please describe your experience:'
    ]
    
    for question in project_questions:
        doc.add_paragraph(question, style='List Number')
        # Add space for answer
        answer_para = doc.add_paragraph()
        answer_para.add_run('Answer: ')
        answer_para.add_run('_' * 80)
        doc.add_paragraph()  # Extra space
    
    # Learning Goals Section
    doc.add_heading('Learning Goals and Expectations', level=1)
    
    goal_questions = [
        'What specific aspects of LLM-based chip design are you most interested in learning?',
        'Do you have any specific career goals related to hardware design or AI?',
        'What challenges do you expect to face in this course?',
        'How do you prefer to learn new technical concepts? (hands-on, theory-first, examples, etc.)'
    ]
    
    for question in goal_questions:
        doc.add_paragraph(question, style='List Number')
        # Add space for answer
        answer_para = doc.add_paragraph()
        answer_para.add_run('Answer: ')
        answer_para.add_run('_' * 80)
        doc.add_paragraph()  # Extra space
    
    # Additional Information Section
    doc.add_heading('Additional Information', level=1)
    
    doc.add_paragraph(
        'Please provide any additional information about your background, interests, '
        'or concerns that might be relevant to this course:'
    )
    
    # Add space for additional comments
    for _ in range(5):
        doc.add_paragraph('_' * 100)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('Thank you for completing this survey! ')
    footer.add_run(
        'This information will help us provide the best possible learning experience '
        'tailored to your background and goals.'
    ).italic = True
    
    footer2 = doc.add_paragraph()
    footer2.add_run('Course: LLM4ChipDesign - Generative AI for Chip Design\n')
    footer2.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y")}\n')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def main():
    """Main function to generate and save the survey form."""
    print("Generating LLM4ChipDesign Course Prerequisite Survey Form...")
    
    # Create the survey document
    survey_doc = create_survey_form()
    
    # Save the document
    output_path = '/home/runner/work/LLM4ChipDesign/LLM4ChipDesign/LLM4ChipDesign_Prerequisite_Survey.docx'
    survey_doc.save(output_path)
    
    print(f"Survey form successfully generated and saved as: {output_path}")
    print("\nThe survey form includes the following sections:")
    print("1. Student Information")
    print("2. Prerequisite Course Assessment (15 key subject areas)")
    print("3. Specific Course Experience")
    print("4. Tools and Software Experience")
    print("5. Project Experience")
    print("6. Learning Goals and Expectations")
    print("7. Additional Information")
    
    # Verify file was created
    if os.path.exists(output_path):
        file_size = os.path.getsize(output_path)
        print(f"\nFile created successfully! Size: {file_size} bytes")
    else:
        print("\nError: File was not created!")
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())